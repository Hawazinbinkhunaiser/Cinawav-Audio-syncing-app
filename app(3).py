
import webbrowser
import streamlit as st
import librosa
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt
import librosa.display
import tempfile
from docx.shared import Inches
import subprocess
import sys
import os 
import psutil

def is_streamlit_running():
    """Check if Streamlit is already running to avoid multiple instances."""
    for process in psutil.process_iter(['name', 'cmdline']):
        if process.info['name'] == 'python.exe' or process.info['name'] == 'python3':
            if any('streamlit' in cmd for cmd in process.info['cmdline']):
                return True
    return False

if __name__ == "__main__":
    if not is_streamlit_running():
        script_path = os.path.abspath(__file__)  # Path to the current script
        subprocess.run([sys.executable, "-m", "streamlit", "run", script_path])
    else:
        print("Streamlit is already running.")
# Function to find the offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

# Function to find the offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

# Process segment data function
def process_segment_data(args):
    interval, master, sample, sr_master, segment_length = args
    start = interval * sr_master
    end = start + segment_length * sr_master  # Segment of defined length
    if end <= len(master) and end <= len(sample):
        master_segment = master[start:end]
        sample_segment = sample[start:end]
        offset = find_offset(master_segment, sample_segment, sr_master)
        return (interval // 60, offset)
    return None

# Function to generate DOCX with results
def generate_docx(results, intervals, dropouts, plots, folder_name):
    doc = Document()
    
    # Add introductory text with today's date and folder name
    specified_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"{folder_name} - {specified_date}", 0)

    # List the names of the devices - extract base names from paths
    device_names = [os.path.basename(name).replace('.wav', '') for name in results.keys()]
    doc.add_paragraph(f"Devices compared: {', '.join(device_names)}\n")

    # Add a table with results
    table = doc.add_table(rows=1, cols=len(device_names) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time (mins)'
    for i, device_name in enumerate(device_names):
        hdr_cells[i + 1].text = device_name

    # Fill the table with intervals and results
    for interval in intervals:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{interval // 60} mins"  # Convert seconds to minutes
        for i, (sample_path, sample_name) in enumerate(zip(results.keys(), device_names)):
            result = next((offset for (intv, offset) in results[sample_path] if intv == interval // 60), None)
            row_cells[i + 1].text = f"{result:.2f} ms" if result is not None else "N/A"

    # Add a section for dropouts
    doc.add_heading("Detected Dropouts", 1)
    for device_path, device_dropouts in dropouts.items():
        device_name = os.path.basename(device_path).replace('.wav', '')
        doc.add_paragraph(f"Dropouts for {device_name}:")
        for dropout in device_dropouts:
            start, end, duration_ms = dropout
            doc.add_paragraph(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
        
        # Add the plot
        doc.add_picture(plots[device_path], width=Inches(6))

    # Add a comments section
    doc.add_paragraph("\nResults and Comments:\n")

    return doc
# Function to format time in HH:MM:SS
def format_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f'{hours:02d}:{minutes:02d}:{secs:06.3f}'

# Function to detect audio dropouts
def detect_dropouts(file_path, dropout_db_threshold=-20, min_duration_ms=100):
    # Load the audio file
    y, sr = librosa.load(file_path, sr=None)

    # Improved time resolution by reducing hop length
    hop_length = 256  # Reduced hop length for better time resolution
    frame_length = hop_length / sr * 1000  # ms per frame

    # Convert the signal to decibels
    rms = librosa.feature.rms(y=y, frame_length=hop_length, hop_length=hop_length)
    rms_db = librosa.power_to_db(rms, ref=np.max)

    # Threshold to find dropouts (segments below the dropout_db_threshold)
    dropout_frames = rms_db[0] < dropout_db_threshold

    # Detect contiguous frames of dropouts lasting at least min_duration_ms
    min_frames = int(min_duration_ms / frame_length)
    dropouts = []
    start = None

    for i, is_dropout in enumerate(dropout_frames):
        if is_dropout and start is None:
            start = i  # Start of a dropout
        elif not is_dropout and start is not None:
            if i - start >= min_frames:
                start_time = start * hop_length / sr
                end_time = i * hop_length / sr
                duration_ms = (end_time - start_time) * 1000  # Convert duration to milliseconds
                dropouts.append((start_time, end_time, duration_ms))
            start = None

    # Handle the case where dropout extends to the end of the file
    if start is not None and len(dropout_frames) - start >= min_frames:
        start_time = start * hop_length / sr
        end_time = len(dropout_frames) * hop_length / sr
        duration_ms = (end_time - start_time) * 1000
        dropouts.append((start_time, end_time, duration_ms))

    return dropouts

# Function to plot waveform with dropouts
def plot_waveform_with_dropouts(y, sr, dropouts, file_name):
    plt.figure(figsize=(12, 6))
    librosa.display.waveshow(y, sr=sr, alpha=0.6)
    plt.title('Waveform with Detected Dropouts')
    plt.xlabel('Time (seconds)')
    plt.ylabel('Amplitude')

    # Highlight dropouts
    for dropout in dropouts:
        start_time, end_time, _ = dropout
        plt.axvspan(start_time, end_time, color='red', alpha=0.5, label='Dropout' if 'Dropout' not in plt.gca().get_legend_handles_labels()[1] else "")

    plt.legend()
    plt.savefig(file_name, bbox_inches='tight')
    plt.close()

def get_or_create_extraction_folder(folder_name):
    if 'extraction_folder' not in st.session_state:
        full_path = os.path.abspath(folder_name)
        os.makedirs(full_path, exist_ok=True)
        st.session_state.extraction_folder = full_path
    return st.session_state.extraction_folder

# Streamlit UI
st.title("CineWav Audio processing Hub")

# Step 1: Upload the M4A file
extraction_folder_name = st.text_input("Enter extraction folder name", value="extracted_tracks")
uploaded_file = st.file_uploader("Choose an M4A file", type=["m4a"])

# Channel Mapping for 7.1 Surround Sound
channels = {
    'Front Left (FL)': 'FL',
    'Front Right (FR)': 'FR',
    'Center (FC)': 'FC',
    'Subwoofer (LFE)': 'LFE',
    'Back Left (BL)': 'BL',
    'Back Right (BR)': 'BR',
    'Side Left (SL)': 'SL',
    'Side Right (SR)': 'SR'
}


if uploaded_file is not None:
    # Get or create the extraction folder
    extraction_folder = get_or_create_extraction_folder(extraction_folder_name)
    st.write(f"Using extraction folder: {extraction_folder}")

    # Save the uploaded file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=".m4a") as temp_file:
        temp_file.write(uploaded_file.getbuffer())
        input_file = temp_file.name

    st.write("File uploaded successfully. Identifying and extracting audio channels...")

    # Step 2: Extract each channel using FFmpeg pan filter
    output_files = []
    track_names = {}
    for name, channel in channels.items():
        track_name = st.text_input(f"Enter track name for {name}", value=name)
        output_file_name = f"{track_name}.wav"
        output_file = os.path.join(extraction_folder, output_file_name)
        if not os.path.exists(output_file):  # Only extract if the file doesn't exist
            command = f'ffmpeg -y -i "{input_file}" -filter_complex "pan=mono|c0={channel}" "{output_file}"'
            subprocess.run(command, shell=True)
            st.write(f"Extracted {name} to {output_file_name}")  # Show only the file name
        output_files.append(output_file)
        track_names[output_file_name] = track_name  # Map file name to custom track name
    # Step 3: Provide download links for each extracted channel
    for output_file in output_files:
        file_name = os.path.basename(output_file)
        with open(output_file, "rb") as f:
            st.download_button(label=f"Download {file_name}", data=f, file_name=file_name, mime="audio/wav")
    # Step 4: Audio Sync Offset Finder and Dropout Detection
    st.subheader("Audio Sync Offset Finder and Dropout Detection")
    st.write("Select a master track and one or more sample tracks to compare.")
    # File selection
    master_file = st.selectbox("Select Master Track", output_files, format_func=lambda x: os.path.basename(x))
    sample_files = st.multiselect("Select Sample Tracks", output_files, default=[output_file for output_file in output_files if output_file != master_file], format_func=lambda x: os.path.basename(x))

# Sampling rate and segment settings
    low_sr = st.slider("Select lower sampling rate for faster processing", 4000, 16000, 4000)
    segment_length = st.slider("Segment length (seconds)", 2, 120, 10)
    intervals = st.multiselect("Select intervals (in seconds)", options=[60, 900, 1800, 2700, 3600, 4500, 5400, 6300], default=[60, 900, 1800, 2700, 3600])

    # Add a "Process" button
    if st.button("Process"):
        if master_file and sample_files:
            st.write("Processing started...")

            # Load the master track
            master, sr_master = librosa.load(master_file, sr=low_sr)

            all_results = {}
            all_dropouts = {}
            all_plots = {}

            for sample_file in sample_files:
                sample, sr_sample = librosa.load(sample_file, sr=low_sr)

                # Resample if the sampling rates do not match
                if sr_master != sr_sample:
                    sample = librosa.resample(sample, sr_sample, sr_master)

                args = [(interval, master, sample, sr_master, segment_length) for interval in intervals]

                with ProcessPoolExecutor() as executor:
                    results = list(filter(None, executor.map(process_segment_data, args)))

                all_results[sample_file] = results

                # Detect dropouts in the sample track
                dropouts = detect_dropouts(sample_file)
                all_dropouts[sample_file] = dropouts

                # Plot waveform with dropouts
                plot_file_name = f"{sample_file}_plot.png"
                plot_waveform_with_dropouts(sample, sr_master, dropouts, plot_file_name)
                all_plots[sample_file] = plot_file_name

            st.write("Processing completed.")
            
            # Display results
            for sample_name, results in all_results.items():
                file_name = os.path.basename(sample_name)
                track_name = track_names[file_name]  # Get custom track name from file name
                st.subheader(f"Results for {track_name}:")
                for interval, offset in results:
                     st.write(f"At {interval // 60} mins: Offset = {offset:.2f} ms")  # Update interval display if necessary

            # Display dropouts
            for sample_name, dropouts in all_dropouts.items():
                file_name = os.path.basename(sample_name)
                st.subheader(f"Detected Dropouts for {file_name}:")
                if dropouts:
                    for dropout in dropouts:
                        start, end, duration_ms = dropout
                        st.write(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
                else:
                    st.write("No significant dropouts detected.")

            # Generate DOCX and provide download option
            doc = generate_docx(all_results, intervals, all_dropouts, all_plots,extraction_folder)
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            # Generate DOCX and provide download option
            doc = generate_docx(all_results, intervals, all_dropouts, all_plots, extraction_folder_name)
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            st.download_button("Download Results as DOCX", data=doc_buffer.getvalue(), file_name=f"{extraction_folder_name}_{datetime.now().strftime('%Y-%m-%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.warning("Please select a master track and at least one sample track to begin processing.")
else:
    st.warning("Please upload an M4A file to begin.")