import subprocess
import streamlit as st
import librosa
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt
import librosa.display
import tempfile
from docx.shared import Inches

# Function to extract channels from M4A file
def extract_channels(input_file, channels):
    output_files = []
    for name, channel in channels.items():
        output_file = f"{name.replace(' ', '_').lower()}.wav"
        command = f'ffmpeg -i "{input_file}" -filter_complex "pan=mono|c0={channel}" "{output_file}"'
        subprocess.run(command, shell=True)
        output_files.append(output_file)
    return output_files

# Function to find offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

# Function to process segment data
def process_segment_data(args):
    interval, master, sample, sr_master, segment_length = args
    start = interval * sr_master
    end = start + segment_length * sr_master
    if end <= len(master) and end <= len(sample):
        master_segment = master[start:end]
        sample_segment = sample[start:end]
        offset = find_offset(master_segment, sample_segment, sr_master)
        return (interval // 60, offset)
    return None

# Function to detect dropouts
def detect_dropouts(file_path, dropout_db_threshold=-20, min_duration_ms=100):
    y, sr = librosa.load(file_path, sr=None)
    hop_length = 256
    frame_length = hop_length / sr * 1000
    rms = librosa.feature.rms(y=y, frame_length=hop_length, hop_length=hop_length)
    rms_db = librosa.power_to_db(rms, ref=np.max)
    dropout_frames = rms_db[0] < dropout_db_threshold
    min_frames = int(min_duration_ms / frame_length)
    dropouts = []
    start = None
    for i, is_dropout in enumerate(dropout_frames):
        if is_dropout and start is None:
            start = i
        elif not is_dropout and start is not None:
            if i - start >= min_frames:
                start_time = start * hop_length / sr
                end_time = i * hop_length / sr
                duration_ms = (end_time - start_time) * 1000
                dropouts.append((start_time, end_time, duration_ms))
            start = None
    if start is not None and len(dropout_frames) - start >= min_frames:
        start_time = start * hop_length / sr
        end_time = len(dropout_frames) * hop_length / sr
        duration_ms = (end_time - start_time) * 1000
        dropouts.append((start_time, end_time, duration_ms))
    return dropouts

# Function to plot waveform with dropouts
def plot_waveform_with_dropouts(y, sr, dropouts, file_name):
    plt.figure(figsize=(12, 6))
    librosa.display.waveshow(y, sr=sr, alpha=0.6)
    plt.title('Waveform with Detected Dropouts')
    plt.xlabel('Time (seconds)')
    plt.ylabel('Amplitude')
    for dropout in dropouts:
        start_time, end_time, _ = dropout
        plt.axvspan(start_time, end_time, color='red', alpha=0.5)
    plt.legend()
    plt.savefig(file_name, bbox_inches='tight')
    plt.close()

# Function to generate DOCX
def generate_docx(results, intervals, dropouts, plots):
    doc = Document()
    specified_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Audio Sync Results - {specified_date}", 0)
    device_names = [name for name in results.keys()]
    doc.add_paragraph(f"Devices compared: {', '.join(device_names)}\n")
    table = doc.add_table(rows=1, cols=len(device_names) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time (mins)'
    for i, device_name in enumerate(device_names):
        hdr_cells[i + 1].text = device_name
    for interval in intervals:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{interval // 60} mins"
        for i, sample_name in enumerate(device_names):
            result = next((offset for (intv, offset) in results[sample_name] if intv == interval // 60), None)
            row_cells[i + 1].text = f"{result:.2f} ms" if result is not None else "N/A"
    doc.add_heading("Detected Dropouts", 1)
    for device_name, device_dropouts in dropouts.items():
        doc.add_paragraph(f"Dropouts for {device_name}:")
        for dropout in device_dropouts:
            start, end, duration_ms = dropout
            doc.add_paragraph(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
        doc.add_picture(plots[device_name], width=Inches(6))
    return doc

# Helper function to format time
def format_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f'{hours:02d}:{minutes:02d}:{secs:06.3f}'

# Streamlit app setup
st.title("7.1 Audio Channel Extractor & Sync Processor")

# Step 1: Upload M4A file and extract WAV files
st.subheader("Step 1: Upload your M4A file for channel extraction")
uploaded_file = st.file_uploader("Choose an M4A file", type=["m4a"])
channels = {
    'Front Left (FL)': 'FL',
    'Front Right (FR)': 'FR',
    'Center (FC)': 'FC',
    'Subwoofer (LFE)': 'LFE',
    'Back Left (BL)': 'BL',
    'Back Right (BR)': 'BR',
    'Side Left (SL)': 'SL',
    'Side Right (SR)': 'SR'
}

if uploaded_file:
    with open("input_file.m4a", "wb") as f:
        f.write(uploaded_file.getbuffer())
    output_files = extract_channels("input_file.m4a", channels)
    st.write("Extraction complete. Download your files below:")
    for output_file in output_files:
        with open(output_file, "rb") as f:
            st.download_button(label=f"Download {output_file}", data=f, file_name=output_file, mime="audio/wav")

# Step 2: Upload master and sample WAV files
st.subheader("Step 2: Upload Master and Sample Tracks for Sync & Dropout Detection")
master_file = st.file_uploader("Upload Master Track", type=["wav"])
sample_files = st.file_uploader("Upload Sample Tracks", type=["wav"], accept_multiple_files=True)

# Process settings
low_sr = st.slider("Select lower sampling rate for faster processing", 4000, 16000, 4000)
segment_length = st.slider("Segment length (seconds)", 2, 120, 10)
intervals = st.multiselect("Select intervals (in seconds)", options=[60, 900, 1800, 2700, 3600, 4500, 5400, 6300], default=[60, 900, 1800, 2700, 3600])

# Step 3: Process tracks
if st.button("Process"):
    if master_file and sample_files:
        st.write("Processing started...")
        with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
            tmp.write(master_file.getvalue())
            tmp.flush()
            master, sr_master = librosa.load(tmp.name, sr=low_sr)

        all_results = {}
        all_dropouts = {}
        all_plots = {}

        for sample_file in sample_files:
            with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
                tmp.write(sample_file.getvalue())
                tmp.flush()
                sample, sr_sample = librosa.load(tmp.name, sr=low_sr)

            if sr_master != sr_sample:
                sample = librosa.resample(sample, sr_sample, sr_master)

            args = [(interval, master, sample, sr_master, segment_length) for interval in intervals]

            with ProcessPoolExecutor() as executor:
                results = list(filter(None, executor.map(process_segment_data, args)))

            all_results[sample_file.name] = results

            dropouts = detect_dropouts(tmp.name)
            all_dropouts[sample_file.name] = dropouts

            plot_file_name = f"{sample_file.name}_plot.png"
            plot_waveform_with_dropouts(sample, sr_master, dropouts, plot_file_name)
            all_plots[sample_file.name] = plot_file_name

        doc = generate_docx(all_results, intervals, all_dropouts, all_plots)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button(label="Download DOCX Report", data=doc_io, file_name="sync_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.write("Processing complete!")
    else:
        st.error("Please upload both master and sample files.")
