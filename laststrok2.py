import streamlit as st
import librosa
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt
import librosa.display
import tempfile
from docx.shared import Inches

# Function to find the offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

# Move the process_segment_data function to the top level
def process_segment_data(args):
    interval, master, sample, sr_master, segment_length = args
    start = interval * sr_master
    end = start + segment_length * sr_master  # Segment of defined length
    if end <= len(master) and end <= len(sample):
        master_segment = master[start:end]
        sample_segment = sample[start:end]
        offset = find_offset(master_segment, sample_segment, sr_master)
        return (interval // 60, offset)
    return None

# Function to generate DOCX with results
def generate_docx(results, intervals, dropouts, plots):
    doc = Document()

    # Add introductory text with today's date
    specified_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Audio Sync Results - {specified_date}", 0)

    # List the names of the devices
    device_names = [name for name in results.keys()]
    doc.add_paragraph(f"Devices compared: {', '.join(device_names)}\n")

    # Add a table with results
    table = doc.add_table(rows=1, cols=len(device_names) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time (mins)'
    for i, device_name in enumerate(device_names):
        hdr_cells[i + 1].text = device_name

    # Fill the table with intervals and results
    for interval in intervals:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{interval // 60} mins"  # Convert seconds to minutes
        for i, sample_name in enumerate(device_names):
            result = next((offset for (intv, offset) in results[sample_name] if intv == interval // 60), None)
            row_cells[i + 1].text = f"{result:.2f} ms" if result is not None else "N/A"

    # Add a section for dropouts
    doc.add_heading("Detected Dropouts", 1)
    for device_name, device_dropouts in dropouts.items():
        doc.add_paragraph(f"Dropouts for {device_name}:")
        for dropout in device_dropouts:
            start, end, duration_ms = dropout
            doc.add_paragraph(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
        
        # Add the plot
        doc.add_picture(plots[device_name], width=Inches(6))

    # Add a comments section
    doc.add_paragraph("\nResults and Comments:\n")
    
    return doc

# Function to format time in HH:MM:SS
def format_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f'{hours:02d}:{minutes:02d}:{secs:06.3f}'

# Function to detect audio dropouts
def detect_dropouts(file_path, dropout_db_threshold=-20, min_duration_ms=100):
    # Load the audio file
    y, sr = librosa.load(file_path, sr=None)

    # Improved time resolution by reducing hop length
    hop_length = 256  # Reduced hop length for better time resolution
    frame_length = hop_length / sr * 1000  # ms per frame

    # Convert the signal to decibels
    rms = librosa.feature.rms(y=y, frame_length=hop_length, hop_length=hop_length)
    rms_db = librosa.power_to_db(rms, ref=np.max)

    # Threshold to find dropouts (segments below the dropout_db_threshold)
    dropout_frames = rms_db[0] < dropout_db_threshold

    # Detect contiguous frames of dropouts lasting at least min_duration_ms
    min_frames = int(min_duration_ms / frame_length)
    dropouts = []
    start = None

    for i, is_dropout in enumerate(dropout_frames):
        if is_dropout and start is None:
            start = i  # Start of a dropout
        elif not is_dropout and start is not None:
            if i - start >= min_frames:
                start_time = start * hop_length / sr
                end_time = i * hop_length / sr
                duration_ms = (end_time - start_time) * 1000  # Convert duration to milliseconds
                dropouts.append((start_time, end_time, duration_ms))
            start = None

    # Handle the case where dropout extends to the end of the file
    if start is not None and len(dropout_frames) - start >= min_frames:
        start_time = start * hop_length / sr
        end_time = len(dropout_frames) * hop_length / sr
        duration_ms = (end_time - start_time) * 1000
        dropouts.append((start_time, end_time, duration_ms))

    return dropouts

# Function to plot waveform with dropouts
def plot_waveform_with_dropouts(y, sr, dropouts, file_name):
    plt.figure(figsize=(12, 6))
    librosa.display.waveshow(y, sr=sr, alpha=0.6)
    plt.title('Waveform with Detected Dropouts')
    plt.xlabel('Time (seconds)')
    plt.ylabel('Amplitude')

    # Highlight dropouts
    for dropout in dropouts:
        start_time, end_time, _ = dropout
        plt.axvspan(start_time, end_time, color='red', alpha=0.5, label='Dropout' if 'Dropout' not in plt.gca().get_legend_handles_labels()[1] else "")
    
    plt.legend()
    plt.savefig(file_name, bbox_inches='tight')
    plt.close()

# Streamlit UI
st.title("Audio Sync Offset Finder and Dropout Detection")
st.write("Upload a master track and one or more sample tracks to compare.")

# File upload
master_file = st.file_uploader("Upload Master Track", type=["wav"])
sample_files = st.file_uploader("Upload Sample Tracks", type=["wav"], accept_multiple_files=True)

# Sampling rate and segment settings
low_sr = st.slider("Select lower sampling rate for faster processing", 4000, 16000, 4000)
segment_length = st.slider("Segment length (seconds)", 2, 120, 10)
intervals = st.multiselect("Select intervals (in seconds)", options=[60, 900, 1800, 2700, 3600, 4500, 5400, 6300], default=[60, 900, 1800, 2700, 3600])

# Add a "Process" button
if st.button("Process"):
    if master_file and sample_files:
        st.write("Processing started...")

        # Load the master track
        with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
            tmp.write(master_file.getvalue())
            tmp.flush()
            master, sr_master = librosa.load(tmp.name, sr=low_sr)

        all_results = {}
        all_dropouts = {}
        all_plots = {}

        for sample_file in sample_files:
            with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
                tmp.write(sample_file.getvalue())
                tmp.flush()
                sample, sr_sample = librosa.load(tmp.name, sr=low_sr)

            # Resample if the sampling rates do not match
            if sr_master != sr_sample:
                sample = librosa.resample(sample, sr_sample, sr_master)

            args = [(interval, master, sample, sr_master, segment_length) for interval in intervals]

            with ProcessPoolExecutor() as executor:
                results = list(filter(None, executor.map(process_segment_data, args)))

            all_results[sample_file.name] = results

            # Detect dropouts in the sample track
            with tempfile.NamedTemporaryFile(suffix=".wav") as tmp:
                tmp.write(sample_file.getvalue())
                tmp.flush()
                dropouts = detect_dropouts(tmp.name)
            all_dropouts[sample_file.name] = dropouts

            # Plot waveform with dropouts
            plot_file_name = f"{sample_file.name}_plot.png"
            plot_waveform_with_dropouts(sample, sr_master, dropouts, plot_file_name)
            all_plots[sample_file.name] = plot_file_name

        st.write("Processing completed.")
        
        # Display results
        for sample_name, results in all_results.items():
            st.subheader(f"Results for {sample_name}:")
            for interval, offset in results:
                st.write(f"At {interval} mins: Offset = {offset:.2f} ms")
        
        # Display dropouts
        for sample_name, dropouts in all_dropouts.items():
            st.subheader(f"Detected Dropouts for {sample_name}:")
            if dropouts:
                for dropout in dropouts:
                    start, end, duration_ms = dropout
                    st.write(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
            else:
                st.write("No significant dropouts detected.")

        # Generate DOCX and provide download option
        doc = generate_docx(all_results, intervals, all_dropouts, all_plots)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button("Download Results as DOCX", data=doc_buffer.getvalue(), file_name="audio_sync_results.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Please upload a master track and at least one sample track to begin processing.")
