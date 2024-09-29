import streamlit as st
import librosa
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from datetime import datetime
from io import BytesIO

# Function to find the offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

# Move the process_segment_data function to the top level
def process_segment_data(args):
    interval, master, sample, sr_master, segment_length = args
    start = interval * sr_master
    end = start + segment_length * sr_master  # Segment of defined length
    if end <= len(master) and end <= len(sample):
        master_segment = master[start:end]
        sample_segment = sample[start:end]
        offset = find_offset(master_segment, sample_segment, sr_master)
        return (interval // 60, offset)
    return None

# Function to generate DOCX with results
def generate_docx(results, intervals):
    doc = Document()

    # Add introductory text with today's date
    specified_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Audio Sync Results - {specified_date}", 0)

    # List the names of the devices
    device_names = [name for name in results.keys()]
    doc.add_paragraph(f"Devices compared: {', '.join(device_names)}\n")

    # Add a table with results
    table = doc.add_table(rows=1, cols=len(device_names) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time (mins)'
    for i, device_name in enumerate(device_names):
        hdr_cells[i + 1].text = device_name

    # Fill the table with intervals and results
    for interval in intervals:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{interval // 60} mins"  # Convert seconds to minutes
        for i, sample_name in enumerate(device_names):
            result = next((offset for (intv, offset) in results[sample_name] if intv == interval // 60), None)
            row_cells[i + 1].text = f"{result:.2f} ms" if result is not None else "N/A"

    # Add a comments section
    doc.add_paragraph("\nResults and Comments:\n")
    
    return doc

# Streamlit UI
st.title("Audio Sync Offset Finder")
st.write("Upload a master track and one or more sample tracks to compare.")

# File upload
master_file = st.file_uploader("Upload Master Track", type=["wav"])
sample_files = st.file_uploader("Upload Sample Tracks", type=["wav"], accept_multiple_files=True)

# Sampling rate and segment settings
low_sr = st.slider("Select lower sampling rate for faster processing", 4000, 16000, 4000)
segment_length = st.slider("Segment length (seconds)", 2, 120, 10)
intervals = st.multiselect("Select intervals (in seconds)", options=[60, 900, 1800, 2700, 3600, 4500, 5400, 6300], default=[60, 900, 1800, 2700, 3600])

# Add a "Process" button
if st.button("Process"):
    if master_file and sample_files:
        st.write("Processing started...")

        # Load the master track
        master, sr_master = librosa.load(master_file, sr=low_sr)

        all_results = {}

        for sample_file in sample_files:
            sample, sr_sample = librosa.load(sample_file, sr=low_sr)

            # Resample if the sampling rates do not match
            if sr_master != sr_sample:
                sample = librosa.resample(sample, sr_sample, sr_master)

            args = [(interval, master, sample, sr_master, segment_length) for interval in intervals]

            with ProcessPoolExecutor() as executor:
                results = list(filter(None, executor.map(process_segment_data, args)))

            all_results[sample_file.name] = results

        st.write("Processing completed.")
        
        # Display results
        for sample_name, results in all_results.items():
            st.subheader(f"Results for {sample_name}:")
            for interval, offset in results:
                st.write(f"At {interval} mins: Offset = {offset:.2f} ms")
        
        # Generate DOCX and provide download option
        doc = generate_docx(all_results, intervals)
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button("Download Results as DOCX", data=doc_buffer.getvalue(), file_name="audio_sync_results.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Please upload a master track and at least one sample track to begin processing.")
