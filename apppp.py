import streamlit as st
import librosa
import librosa.display
import numpy as np
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from datetime import datetime
from io import BytesIO
import matplotlib.pyplot as plt

# Function to find the offset between master and sample segments
def find_offset(master_segment, sample_segment, sr):
    correlation = np.correlate(master_segment, sample_segment, mode='full')
    max_corr_index = np.argmax(correlation)
    offset_samples = max_corr_index - len(sample_segment) + 1
    offset_ms = (offset_samples / sr) * 1000  # Convert to milliseconds
    return offset_ms

def process_segment_data(interval, master_path, sample_path, sr, segment_length):
    # Load the audio data within the process
    master, _ = librosa.load(master_path, sr=None)
    sample, _ = librosa.load(sample_path, sr=None)

    start = interval * sr
    end = start + segment_length * sr  # Segment of defined length
    if end <= len(master) and end <= len(sample):
        master_segment = master[start:end]
        sample_segment = sample[start:end]
        offset = find_offset(master_segment, sample_segment, sr)
        return (interval // 60, offset)
    return None

# Function to format time in HH:MM:SS
def format_time(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f'{hours:02d}:{minutes:02d}:{secs:06.3f}'

# Function to detect audio dropouts
def detect_dropouts(file_data, dropout_db_threshold=-20, min_duration_ms=100):
    # Load the audio file
    y, sr = librosa.load(file_data, sr=None)

    # Improved time resolution by reducing hop length
    hop_length = 256  # Reduced hop length for better time resolution
    frame_length = hop_length / sr * 1000  # ms per frame

    # Convert the signal to decibels
    rms = librosa.feature.rms(y=y, frame_length=hop_length, hop_length=hop_length)
    rms_db = librosa.power_to_db(rms, ref=np.max)

    # Threshold to find dropouts (segments below the dropout_db_threshold)
    dropout_frames = rms_db[0] < dropout_db_threshold

    # Detect contiguous frames of dropouts lasting at least min_duration_ms
    min_frames = int(min_duration_ms / frame_length)
    dropouts = []
    start = None

    for i, is_dropout in enumerate(dropout_frames):
        if is_dropout and start is None:
            start = i  # Start of a dropout
        elif not is_dropout and start is not None:
            if i - start >= min_frames:
                start_time = start * hop_length / sr
                end_time = i * hop_length / sr
                duration_ms = (end_time - start_time) * 1000  # Convert duration to milliseconds
                dropouts.append((start_time, end_time, duration_ms))
            start = None

    # Handle the case where dropout extends to the end of the file
    if start is not None and len(dropout_frames) - start >= min_frames:
        start_time = start * hop_length / sr
        end_time = len(dropout_frames) * hop_length / sr
        duration_ms = (end_time - start_time) * 1000
        dropouts.append((start_time, end_time, duration_ms))

    return dropouts, y, sr

# Function to plot waveform with dropouts
def plot_waveform_with_dropouts(y, sr, dropouts):
    plt.figure(figsize=(12, 6))
    librosa.display.waveshow(y, sr=sr, alpha=0.6)
    plt.title('Waveform with Detected Dropouts')
    plt.xlabel('Time (seconds)')
    plt.ylabel('Amplitude')

    # Highlight dropouts
    for dropout in dropouts:
        start_time, end_time, _ = dropout
        plt.axvspan(start_time, end_time, color='red', alpha=0.5, label='Dropout' if 'Dropout' not in plt.gca().get_legend_handles_labels()[1] else "")
    
    plt.legend()
    st.pyplot(plt)

# Function to generate DOCX with results
def generate_docx(sync_results, dropout_results, intervals):
    doc = Document()

    # Add introductory text with today's date
    specified_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Audio Sync and Dropout Results - {specified_date}", 0)

    # List the names of the devices
    device_names = [name for name in sync_results.keys()]
    doc.add_paragraph(f"Devices compared: {', '.join(device_names)}\n")

    # Add Sync results table
    doc.add_heading("Sync Results", level=1)
    table = doc.add_table(rows=1, cols=len(device_names) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Time (mins)'
    for i, device_name in enumerate(device_names):
        hdr_cells[i + 1].text = device_name

    for interval in intervals:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{interval // 60} mins"
        for i, sample_name in enumerate(device_names):
            result = next((offset for (intv, offset) in sync_results[sample_name] if intv == interval // 60), None)
            row_cells[i + 1].text = f"{result:.2f} ms" if result is not None else "N/A"

    # Add Dropout results
    doc.add_heading("Dropout Detection", level=1)
    for device_name, dropouts in dropout_results.items():
        doc.add_paragraph(f"Dropouts for {device_name}:")
        for dropout in dropouts:
            start, end, duration_ms = dropout
            doc.add_paragraph(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
    
    return doc

# Streamlit UI
st.title("Audio Sync and Dropout Detection")
st.write("Upload a master track and one or more sample tracks to compare sync offsets and detect audio dropouts.")

# File upload
master_file = st.file_uploader("Upload Master Track", type=["wav"])
sample_files = st.file_uploader("Upload Sample Tracks", type=["wav"], accept_multiple_files=True)

# Sampling rate and segment settings
low_sr = st.slider("Select lower sampling rate for faster processing", 4000, 16000, 4000)
segment_length = st.slider("Segment length (seconds)", 2, 120, 10)
intervals = st.multiselect("Select intervals (in seconds)", options=[60, 900, 1800, 2700, 3600, 4500, 5400, 6300], default=[60, 900, 1800, 2700, 3600])

# Add a "Process" button
if st.button("Process"):
    if master_file and sample_files:
        st.write("Processing started...")

        # Load the master track
        master_path = master_file.name  # Get the path from the uploaded file
        master, sr_master = librosa.load(master_file, sr=low_sr)

        all_sync_results = {}
        all_dropout_results = {}

        for sample_file in sample_files:
            sample_path = sample_file.name  # Get the path from the uploaded file
            sample, sr_sample = librosa.load(sample_file, sr=low_sr)

            # Resample if the sampling rates do not match
            if sr_master != sr_sample:
                sample = librosa.resample(sample, sr_sample, sr_master)

            # Process sync offsets
            args = [(interval, master_path, sample_path, sr_master, segment_length) for interval in intervals]

            with ProcessPoolExecutor() as executor:
                sync_results = list(filter(None, executor.map(lambda x: process_segment_data(*x), args)))

            all_sync_results[sample_file.name] = sync_results

            # Process dropout detection
            dropouts, y, sr = detect_dropouts(sample_file)
            all_dropout_results[sample_file.name] = dropouts

            # Display Sync results
            st.subheader(f"Sync Results for {sample_file.name}:")
            for interval, offset in sync_results:
                st.write(f"At {interval} mins: Offset = {offset:.2f} ms")
            
            # Display Dropout results
            st.subheader(f"Dropout Results for {sample_file.name}:")
            if dropouts:
                for dropout in dropouts:
                    start, end, duration_ms = dropout
                    st.write(f"Start: {format_time(start)} | End: {format_time(end)} | Duration: {duration_ms:.0f} ms")
            else:
                st.write("No dropouts detected.")

        # Generate and download DOCX with results
        doc = generate_docx(all_sync_results, all_dropout_results, intervals)
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button("Download Results as DOCX", doc_io, "results.docx")
